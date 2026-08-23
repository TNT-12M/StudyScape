#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档提取与题目解析模块
支持 PDF（含扫描件 OCR）、DOC、DOCX、TXT。

OCR 引擎配置（2025-08-20 定稿）：
    默认: onnxruntime + PP-OCRv5 MOBILE（经实测 2核2G 服务器最优）
        - 初始化 1.1s，单页 OCR ~8s，2页 16s，内存 Δ~38MB
        - 中文识别率 100% 有效（无乱码/问号）
    备选: MNN 引擎 + PP-OCRv5 MOBILE（Linux 下可能更快，需实测）
    兜底: Tesseract CLI（内存 ~50MB，精度较低但稳定）

OCR 优先级（PDF 扫描件）：
    1) pdftotext（文本型 PDF 最快）
    2) 子进程逐页 RapidOCR（默认 onnxruntime+v5-MOBILE）
       每页独立 Python 进程 → 防止单页 OOM 影响全局
    3) 子进程逐页 Tesseract CLI 兜底（内存 ~50MB）
    4) 单进程生成器 RapidOCR（内存充裕时可用）
    5) 单进程 pytesseract 兜底

用法：
    extract.py <文件路径>            # 仅提取纯文本
    extract.py <文件路径> --parse    # 提取文本并尝试解析选择题
    extract.py <文件路径> --ocr       # 强制走 OCR（即使有 pdftotext 结果）

输出（JSON）：
    {
        "success": true,
        "file": "...",
        "type": "pdf",
        "method": "ocr-rapidocr-fitz-subprocess",
        "text": "...",
        "questions": [ ... ]   # 仅 --parse 时存在
    }
"""

# ==========================================================
# 默认 OCR 引擎配置（2025-08-20 定稿）
# ==========================================================
OCR_CONFIG = {
    # 首选引擎：onnxruntime + PP-OCRv5 MOBILE
    "engine": "onnxruntime",       # onnxruntime | mnn (Linux 备选)
    "det_model": "mobile",         # tiny | small | mobile | server
    "rec_model": "mobile",         # tiny | small | mobile | server
    "ocr_version": "v5",           # v5 (默认) | v4 | v6
    "lang_det": "ch",              # ch | en | multi
    "lang_rec": "ch",              # ch | en | ...
    # DPI 配置
    "dpi_rapidocr": 120,           # OCR 分辨率（越高越准但越慢）
    "dpi_tesseract": 120,          # Tesseract 兜底分辨率
    # 内存保护
    "threads": 1,                  # 限制推理线程数，控制内存
    "timeout_per_page": 180,       # 单页超时（秒）
}

# 备选 MNN 配置（Linux 下可能更快）
OCR_CONFIG_MNN = {
    "engine": "mnn",
    "det_model": "mobile",
    "rec_model": "mobile",
    "ocr_version": "v5",
    "lang_det": "ch",
    "lang_rec": "ch",
}

import io
import json
import os
import re
import subprocess
import sys
import tempfile

# 在模块导入早期（RapidOCR import 之前）就把 RapidOCR 家族 logger 静音到 WARNING，
# 防止 "Using engine_name" / "File exists and is valid" / "Using xxx.onnx" 等
# INFO 日志通过 stderr 混进 shell_exec 的 2>&1 合并输出，导致 PHP json_decode 失败。
import logging as _logging_root

# 1. 强制基础配置在任何 import rapidocr 之前生效：默认 handler 只接收 WARNING 及以上，输出到 stderr
_logging_root.basicConfig(level=_logging_root.WARNING, force=True, stream=sys.stderr)

# 2. RapidOCR 相关 logger 单独静音并阻断向上传播
for _ln in ('RapidOCR', 'rapidocr', 'rapidocr_onnxruntime'):
    _lg = _logging_root.getLogger(_ln)
    _lg.setLevel(_logging_root.WARNING)
    _lg.propagate = False  # 关键：不向父 logger 传播，完全阻断 INFO 流
    _lg.handlers = [_h for _h in _lg.handlers if not isinstance(_h, _logging_root.StreamHandler)]


# ==========================================================
# 工具：解码 / 命令执行
# ==========================================================

def smart_decode(raw: bytes) -> str:
    """
    智能解码 CLI 命令输出。
    顺序：UTF-8 严格 → 中文 Windows(GB18030/CP936/Big5) → UTF-8 兜底
    antiword / catdoc / pdftotext 在中文 Windows 默认输出 CP936(GBK)，
    而 Python stdout 在 PIPE 下经常被错误按 UTF-8 解码导致全是乱码/问号。
    """
    if raw is None:
        return ""
    if not raw:
        return ""
    # 1. 先严格尝试 UTF-8（antiword -m UTF-8 / pdftotext UTF-8 locale）
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass
    # 2. 中文 Windows (GBK/CP936) — antiword 默认 / catdoc / pdftotext 最常见
    for enc in ("gb18030", "cp936", "big5", "utf-8"):
        try:
            result = raw.decode(enc)
            stripped = result.replace("\r", "").replace("\n", "").replace("\t", "").strip()
            if stripped:
                bad = sum(1 for ch in stripped if ch in ("?",))
                if bad / len(stripped) < 0.4:
                    return result
        except (UnicodeDecodeError, LookupError):
            continue
    # 3. 全部失败，使用 utf-8 忽略错误兜底
    return raw.decode("utf-8", errors="ignore")


def run(cmd, timeout=300, capture_stderr=False):
    """执行外部命令并返回 (returncode, stdout)。"""
    try:
        env = os.environ.copy()
        proc = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE if capture_stderr else subprocess.DEVNULL,
            timeout=timeout, check=False, env=env
        )
        return proc.returncode, smart_decode(proc.stdout)
    except Exception:
        return -1, ""


def is_text_meaningful(text: str, min_chars: int = 50, max_q_ratio: float = 0.4) -> bool:
    """判断文本是否“有足够有效内容”。antiword 不加 -m UTF-8 会输出大量问号。"""
    if not text:
        return False
    cleaned = re.sub(r"\s+", "", text)
    if len(cleaned) < min_chars:
        return False
    # 中文/英文字符比例
    meaningful = sum(1 for c in cleaned if "\u4e00" <= c <= "\u9fff" or c.isascii() and c.isalnum())
    if meaningful == 0:
        return False
    q = cleaned.count("?")
    if len(cleaned) > 0 and q / len(cleaned) > max_q_ratio:
        return False
    return True


# ==========================================================
# 文件类型
# ==========================================================

def file_type(path):
    ext = os.path.splitext(path)[1].lower()
    mapping = {
        ".pdf": "pdf",
        ".doc": "doc",
        ".docx": "docx",
        ".txt": "txt",
        ".md": "txt",
    }
    return mapping.get(ext, ext.lstrip(".") or "unknown")


# ==========================================================
# TXT / DOCX
# ==========================================================

def extract_txt(path):
    # 尝试 UTF-8，再尝试 GB18030
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="gb18030", errors="ignore") as f:
            return f.read()


def extract_docx(path):
    try:
        from docx import Document
    except Exception:
        return "", "docx"
    try:
        doc = Document(path)
    except Exception:
        return "", "docx"
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts), "docx"


# ==========================================================
# DOC（关键：antiword 优先带 -m UTF-8 参数）
# ==========================================================

def extract_doc(path):
    # 1) antiword 带 UTF-8 映射（Git/mingw64 的 antiword 中文默认输出问号，此参数才会得到真中文UTF-8）
    code, out = run(["antiword", "-m", "UTF-8", path])
    if code == 0 and is_text_meaningful(out):
        return out, "antiword-utf8"

    # 2) 退回到 antiword 默认（可能是 GBK/CP936）
    code, out = run(["antiword", path])
    if code == 0 and is_text_meaningful(out):
        return out, "antiword"

    # 3) catdoc
    code, out = run(["catdoc", path])
    if code == 0 and is_text_meaningful(out):
        return out, "catdoc"

    # 4) 实在不行，返回 antiword -m UTF-8 的原始结果（哪怕空也让上层知道是什么 method 导致）
    return out if out else "", "antiword-utf8"


# ==========================================================
# PDF：pdftotext → 文本太少则走 OCR
# 转图优先级：PyMuPDF (fitz，零额外系统依赖：无需 poppler)
#             > pdf2image + poppler (用户已安装 poppler 时)
# OCR 引擎优先级：rapidocr (RapidOCR onnx，用户指定首选)
#             > pdf2image+pytesseract (兼容回退)
# ==========================================================

def _build_rapidocr_params(config=None):
    """根据 OCR_CONFIG 构建 RapidOCR 的 params 字典。
    config 为 None 时使用默认 OCR_CONFIG。
    """
    if config is None:
        config = OCR_CONFIG

    engine_map = {"onnxruntime": "onnxruntime", "mnn": "mnn"}
    model_map = {
        "tiny": "tiny", "small": "small",
        "mobile": "mobile", "server": "server",
    }
    version_map = {"v4": "PPOCRV4", "v5": "PPOCRV5", "v6": "PPOCRV6"}

    engine = engine_map.get(config.get("engine", "onnxruntime"), "onnxruntime")
    det_model = model_map.get(config.get("det_model", "mobile"), "mobile")
    rec_model = model_map.get(config.get("rec_model", "mobile"), "mobile")
    ocr_version = version_map.get(config.get("ocr_version", "v5"), "PPOCRV5")
    lang_det = config.get("lang_det", "ch")
    lang_rec = config.get("lang_rec", "ch")

    try:
        from rapidocr import EngineType, ModelType, OCRVersion, LangDet, LangRec

        engine_type_map = {
            "onnxruntime": EngineType.ONNXRUNTIME,
            "mnn": EngineType.MNN,
        }
        model_type_map = {
            "tiny": ModelType.TINY, "small": ModelType.SMALL,
            "mobile": ModelType.MOBILE, "server": ModelType.SERVER,
        }
        ocr_version_map = {
            "PPOCRV4": OCRVersion.PPOCRV4,
            "PPOCRV5": OCRVersion.PPOCRV5,
            "PPOCRV6": OCRVersion.PPOCRV6,
        }
        lang_det_map = {"ch": LangDet.CH, "en": LangDet.EN, "multi": LangDet.MULTI}
        lang_rec_map = {
            "ch": LangRec.CH, "en": LangRec.EN,
            "chinese_cht": LangRec.CHINESE_CHT,
            "japan": LangRec.JAPAN, "korean": LangRec.KOREAN,
            "latin": LangRec.LATIN, "cyrillic": LangRec.CYRILLIC,
            "arabic": LangRec.ARABIC,
        }

        params = {
            "Det.engine_type": engine_type_map.get(engine, EngineType.ONNXRUNTIME),
            "Det.lang_type": lang_det_map.get(lang_det, LangDet.CH),
            "Det.model_type": model_type_map.get(det_model, ModelType.MOBILE),
            "Det.ocr_version": ocr_version_map.get(ocr_version, OCRVersion.PPOCRV5),
            "Rec.engine_type": engine_type_map.get(engine, EngineType.ONNXRUNTIME),
            "Rec.lang_type": lang_rec_map.get(lang_rec, LangRec.CH),
            "Rec.model_type": model_type_map.get(rec_model, ModelType.MOBILE),
            "Rec.ocr_version": ocr_version_map.get(ocr_version, OCRVersion.PPOCRV5),
        }
        return params
    except Exception:
        return None


def _ocr_rapidocr_images(images, config=None):
    """用 RapidOCR 对图片序列做 OCR。
    images：元素可以是 PIL.Image，或 numpy.ndarray(RGB)，或 图片路径字符串。
    config: OCR 配置字典（None 时使用默认 OCR_CONFIG）
    返回拼接后的纯文本。
    """
    try:
        from rapidocr import RapidOCR
    except Exception:
        try:
            from rapidocr_onnxruntime import RapidOCR
        except Exception:
            return None

    # RapidOCR 的 utils/log.py 在模块首次 import 时就已经创建了
    # Logger(log_level=INFO, logger_name="RapidOCR") 并给它 attach 了一个 StreamHandler。
    # 并且 __init__ 时会通过 logger.setLevel(cfg.Global.log_level.upper()) 把级别重置回 INFO。
    # 所以只能用 monkey patch：在创建 RapidOCR 前后，全局劫持 StreamHandler.emit，
    # 把 logger_name=RapidOCR 的 INFO 记录直接丢弃，完全阻断进入输出流。
    import logging as _logging

    _orig_emit = _logging.StreamHandler.emit
    def _filtered_emit(self_h, record):
        if record.levelno <= _logging.INFO and (record.name == "RapidOCR" or record.name.startswith("rapidocr")):
            return None
        return _orig_emit(self_h, record)
    _logging.StreamHandler.emit = _filtered_emit

    def _silence_rapidocr_loggers():
        for _ln in ('RapidOCR', 'rapidocr', 'rapidocr_onnxruntime'):
            _lg = _logging.getLogger(_ln)
            _lg.setLevel(_logging.WARNING)
            _lg.propagate = False
            _lg.handlers = [_h for _h in _lg.handlers if not isinstance(_h, _logging.StreamHandler)]

    try:
        # 构建配置参数
        params = _build_rapidocr_params(config)
        if params:
            engine = RapidOCR(params=params)
        else:
            engine = RapidOCR()
    except Exception:
        _logging.StreamHandler.emit = _orig_emit
        return None

    # 引擎初始化完成后恢复 handler.emit（后续不再有 RapidOCR 的 INFO 日志污染）
    _logging.StreamHandler.emit = _orig_emit
    _silence_rapidocr_loggers()

    import numpy as np
    import gc
    parts = []
    for img in images:
        arr = None
        if isinstance(img, (str, bytes, os.PathLike)):
            try:
                from PIL import Image
                arr = np.array(Image.open(img).convert("RGB"))
            except Exception:
                parts.append("")
                continue
        else:
            try:
                arr = np.array(img.convert("RGB"))
            except Exception:
                try:
                    arr = np.asarray(img)
                except Exception:
                    arr = img
        try:
            if hasattr(img, 'close'):
                img.close()
        except Exception:
            pass
        try:
            result = engine(arr)
        except Exception:
            parts.append("")
            del arr
            gc.collect()
            continue

        txts = None
        if result is None:
            parts.append("")
            del arr, result
            gc.collect()
            continue
        if hasattr(result, "txts"):
            txts = result.txts
        else:
            try:
                if isinstance(result, tuple) and len(result) >= 1 and isinstance(result[0], list):
                    txts = [item[1] for item in result[0] if isinstance(item, (list, tuple)) and len(item) >= 2]
            except Exception:
                txts = None
        if txts:
            parts.append("\n".join(str(t) for t in txts if t))
        else:
            parts.append("")
        del arr, result
        gc.collect()
    return "\n".join(parts)


def _ocr_tesseract_images(images):
    """回退：pdf2image/pil 转图 + pytesseract OCR。"""
    try:
        import pytesseract
    except Exception:
        return None
    texts = []
    for img in images:
        try:
            # img 可能是 PIL.Image 或路径
            if isinstance(img, (str, bytes, os.PathLike)):
                from PIL import Image
                img = Image.open(img)
            txt = pytesseract.image_to_string(img, lang="chi_sim+eng", config="--psm 6")
        except Exception:
            txt = ""
        texts.append(txt)
    return "\n".join(texts)


def _pdf_render_pages_with_fitz(path, dpi=60):
    """用 PyMuPDF (fitz) 把 PDF 每页转成 PIL.Image(RGB) 生成器（逐页 yield）。
    优点：纯 Python wheel，不需要安装系统 poppler。
    内存优化：生成器而非 list，调用方逐页 OCR 后即可丢弃，避免一次持有所有页（
    9 页 180DPI 的 RGB 图像约 85MB；改 120DPI 后单页约 4MB，且只持有一页）。
    失败时返回 None；返回的生成器可被 for 循环消费一次。
    """
    try:
        import pymupdf as fitz  # PyMuPDF（用新模块名以避免 fitz API 弃用警告污染 stdout/stderr）
    except Exception:
        return None
    try:
        from PIL import Image
    except Exception:
        return None

    try:
        doc = fitz.open(path)
    except Exception:
        return None

    def _iter_pages():
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        with doc:
            for page in doc:
                try:
                    pm = page.get_pixmap(matrix=mat, alpha=False)
                    # pm.samples 是 bytes(RGB)
                    mode = "RGB"
                    if pm.n == 4:
                        mode = "RGBA"
                    img = Image.frombytes(mode, (pm.width, pm.height), pm.samples)
                    if mode == "RGBA":
                        img = img.convert("RGB")
                    yield img
                except Exception:
                    continue

    return _iter_pages()


def _pdf_render_pages_with_pdf2image(path, dpi=120):
    """pdf2image + poppler。仅在 poppler 可用时生效。"""
    try:
        from pdf2image import convert_from_path
    except Exception:
        return None
    try:
        return convert_from_path(path, dpi=dpi)
    except Exception:
        return None


def _ocr_tesseract_pages(pages):
    """回退：pdf2image + pytesseract。"""
    try:
        import pytesseract
    except Exception:
        return None
    texts = []
    for page in pages:
        try:
            txt = pytesseract.image_to_string(page, lang="chi_sim+eng", config="--psm 6")
        except Exception:
            txt = ""
        texts.append(txt)
    return "\n".join(texts)


def _ocr_tesseract_cli(image_path, lang="chi_sim+eng", psm=6, timeout=120):
    """用 tesseract CLI 直接 OCR 单张图片（内存约 50MB，远低于 RapidOCR 的 400MB）。
    在 1.8G 内存机器上当 RapidOCR 被 OOM kill 时的可靠兜底。
    返回纯文本。
    """
    code, out = run(["tesseract", image_path, "stdout",
                     "-l", lang, "--psm", str(psm)], timeout=timeout)
    if code == 0:
        return out
    return ""


def _get_pdf_page_count(path):
    code, out = run(["pdfinfo", path])
    if code == 0:
        m = re.search(r"^Pages:\s+(\d+)", out, re.MULTILINE)
        if m:
            return int(m.group(1))
    try:
        import pymupdf as fitz
        doc = fitz.open(path)
        n = len(doc)
        doc.close()
        return n
    except Exception:
        return 0


def _ocr_pdf_subprocess_per_page(path, dpi_rapidocr=None, dpi_tesseract=None,
                                 timeout_per_page=None, ocr_config=None):
    """子进程逐页 OCR，带 RapidOCR→Tesseract 自动兜底。

    每页先尝试 RapidOCR 子进程（dpi_rapidocr，默认 120，低内存高精度），
    若被 OOM kill（exit 137 / -9）或超时，则同一页用 Tesseract CLI 子进程兜底
    （dpi_tesseract，内存约 50MB，稳定但精度较低）。

    ocr_config: OCR 配置字典（None 时使用 OCR_CONFIG）。
    返回 (拼接文本, 引擎标记) 或 (None, None)。
    """
    if dpi_rapidocr is None:
        dpi_rapidocr = (ocr_config or OCR_CONFIG).get("dpi_rapidocr", 120)
    if dpi_tesseract is None:
        dpi_tesseract = (ocr_config or OCR_CONFIG).get("dpi_tesseract", 120)
    if timeout_per_page is None:
        timeout_per_page = (ocr_config or OCR_CONFIG).get("timeout_per_page", 180)

    py_bin = sys.executable or "python3"
    abs_pdf = os.path.abspath(path)
    abs_self = os.path.abspath(__file__)

    total_pages = _get_pdf_page_count(abs_pdf)
    if total_pages == 0:
        return None, None

    import gc as _gc
    child_env = os.environ.copy()
    child_env["OMP_NUM_THREADS"] = "1"
    child_env["OPENBLAS_NUM_THREADS"] = "1"
    child_env["MKL_NUM_THREADS"] = "1"
    child_env["ONNXRUNTIME_THREAD_COUNT"] = "1"

    # 传递 OCR 配置到子进程（通过环境变量）
    cfg = ocr_config or OCR_CONFIG
    child_env["OCR_ENGINE"] = str(cfg.get("engine", "onnxruntime"))
    child_env["OCR_DET_MODEL"] = str(cfg.get("det_model", "mobile"))
    child_env["OCR_REC_MODEL"] = str(cfg.get("rec_model", "mobile"))
    child_env["OCR_VERSION"] = str(cfg.get("ocr_version", "v5"))

    def _try_engine(page_idx, engine, dpi):
        """对单页启动子进程，返回文本或 None。"""
        try:
            # 传入配置作为额外命令行参数（子进程读取环境变量或命令行参数）
            proc = subprocess.run(
                [py_bin, abs_self, "--ocr-page", abs_pdf,
                 str(page_idx), str(dpi), engine],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=timeout_per_page,
                check=False,
                env=child_env,
            )
        except subprocess.TimeoutExpired:
            return None
        except Exception:
            return None
        if proc.returncode != 0:
            return None
        try:
            out = json.loads(proc.stdout.decode("utf-8", errors="ignore"))
            if out.get("success"):
                return out.get("text", "") or ""
        except Exception:
            pass
        return None

    parts = []
    used_engines = set()
    for page_idx in range(total_pages):
        # 第一轮：RapidOCR（dpi_rapidocr，低内存高精度，可能 OOM）
        text = _try_engine(page_idx, "rapidocr", dpi_rapidocr)
        if text:
            parts.append(text)
            used_engines.add("rapidocr")
        else:
            # 第二轮：Tesseract CLI（dpi_tesseract，低内存稳定兜底）
            text = _try_engine(page_idx, "tesseract", dpi_tesseract)
            if text:
                parts.append(text)
                used_engines.add("tesseract")
            else:
                parts.append("")
        _gc.collect()

    result_text = "\n".join(parts)
    if not result_text.strip():
        return None, None
    if used_engines == {"rapidocr"}:
        tag = "rapidocr"
    elif used_engines == {"tesseract"}:
        tag = "tesseract"
    else:
        tag = "mixed"
    # 在 tag 中加入模型信息（如 rapidocr-ort-v5-mobile）
    engine_short = child_env.get("OCR_ENGINE", "onnxruntime")
    model_short = child_env.get("OCR_DET_MODEL", "mobile")
    tag = f"{tag}-{engine_short[:3]}-v5-{model_short}"
    return result_text, tag


def extract_pdf(path, force_ocr=False):
    """先尝试 pdftotext（文本型 PDF 最快），文本过少则走 OCR。
    OCR 路径（优先级）：
      ① 子进程逐页 OCR（每页独立 Python 进程 + RapidOCR→Tesseract 自动兜底）
         默认配置：onnxruntime + v5-MOBILE（38MB 内存，100% 有效率）
      ② 单进程生成器逐页 OCR（fitz 生成器，适合内存充裕环境）
      ③ pytesseract 回退

    force_ocr=True 时跳过 pdftotext，直接走 OCR（用于测试或强制 OCR 场景）。
    """
    code, out = run(["pdftotext", "-layout", path, "-"])
    if not force_ocr and code == 0 and is_text_meaningful(out, min_chars=50, max_q_ratio=0.4):
        return out, "pdftotext"

    # ========== OCR 路径 ==========
    # 1. 优先：子进程逐页 OCR（默认 onnxruntime+v5-MOBILE，防 OOM）
    sub_text, sub_tag = _ocr_pdf_subprocess_per_page(path)
    if sub_text and is_text_meaningful(sub_text, min_chars=20, max_q_ratio=0.6):
        return sub_text, f"ocr-{sub_tag}-fitz-subprocess"

    # 2. 回退：单进程生成器逐页 OCR（内存够时可用，使用默认 OCR_CONFIG）
    dpi = OCR_CONFIG.get("dpi_rapidocr", 120)
    images_iter = _pdf_render_pages_with_fitz(path, dpi=dpi)
    method_tag = "fitz"
    if images_iter is None:
        images_iter = _pdf_render_pages_with_pdf2image(path, dpi=dpi)
        method_tag = "pdf2image"
    if not images_iter:
        return out if out else "", "pdftotext"

    ocr_text = _ocr_rapidocr_images(images_iter)
    if ocr_text and is_text_meaningful(ocr_text, min_chars=20, max_q_ratio=0.6):
        return ocr_text, f"ocr-rapidocr-{method_tag}"

    # 3. 回退 pytesseract
    if method_tag == "fitz":
        images_iter = _pdf_render_pages_with_fitz(path, dpi=dpi)
    ocr_text2 = _ocr_tesseract_images(images_iter)
    if ocr_text2 and is_text_meaningful(ocr_text2, min_chars=20, max_q_ratio=0.6):
        return ocr_text2, f"ocr-tesseract-{method_tag}"

    # OCR 都失败
    final = sub_text or ocr_text or ocr_text2 or out or ""
    if sub_text:
        m = f"ocr-{sub_tag}-fitz-subprocess"
    elif ocr_text:
        m = f"ocr-rapidocr-{method_tag}"
    elif ocr_text2:
        m = f"ocr-tesseract-{method_tag}"
    else:
        m = f"ocr-{method_tag}"
    return final, m


# ==========================================================
# 题目解析（选择题 + 答案区间自动回填）
# ==========================================================

# 选项标记：A. / A、 / A， / A- / A。 等 OCR 常见变体
OPTION_MARK_RE = re.compile(r"(?<![A-Za-z])([A-H])\s*[.、．，,。\-_:]\s*")
QUESTION_RE = re.compile(r"^(\d{1,3})\s*[.、．，,、]?\s*(.*)$")


def extract_options(line):
    """从一行中提取所有选项，返回 ['A. xxx', 'B. yyy', ...]。"""
    matches = list(OPTION_MARK_RE.finditer(line))
    if not matches:
        return []
    opts = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(line)
        text = line[start:end].strip(" \t,，。;；")
        if text:
            opts.append("%s. %s" % (m.group(1), text))
    return opts


def parse_choice_questions(text):
    """解析选择题（题号 + A-D 选项），返回题目列表。"""
    lines = [l.rstrip() for l in text.split("\n")]
    questions = []
    current = None

    def flush():
        nonlocal current
        if current and len(current["options"]) >= 2 and current["stem"].strip():
            questions.append(current)
        current = None

    for line in lines:
        s = line.strip()
        if not s:
            continue

        # 跳过答案区间行（如 1-5: DBCCB）
        if re.match(r"^\d{1,3}\s*[-~—]\s*\d{1,3}\s*[:：]", s):
            continue

        m = QUESTION_RE.match(s)
        if m:
            rest = m.group(2).strip()
            # 若紧跟选项标记，说明上一题题干为空，忽略该“题号”
            if OPTION_MARK_RE.match(rest):
                if current is not None:
                    current["options"].extend(extract_options(s))
                continue
            flush()
            current = {
                "number": int(m.group(1)),
                "stem": rest,
                "options": [],
            }
            current["options"].extend(extract_options(rest))
            continue

        if current is not None:
            opts = extract_options(s)
            if opts:
                current["options"].extend(opts)
            else:
                current["stem"] += " " + s

    flush()

    # 过滤：题干与选项不能为空，且选项数量合理（2-8）
    result = []
    for q in questions:
        stem = re.sub(r"\s+", " ", q["stem"]).strip()
        if not stem:
            continue
        opts = q["options"]
        if len(opts) < 2 or len(opts) > 8:
            continue
        result.append({
            "number": q["number"],
            "question_type": "single" if len(opts) == 4 else "multiple",
            "content": stem,
            "options": opts,
            "correct_answer": "",
            "points": 2,
        })
    return result


def parse_answer_key(text):
    """解析参考答案区间，如 '1-5: DBCCB'，返回 {题号: 字母}。"""
    answers = {}
    for m in re.finditer(r"(\d{1,3})\s*[-~—]\s*(\d{1,3})\s*[:：]\s*([A-Ha-h]+)", text):
        start = int(m.group(1))
        end = int(m.group(2))
        letters = m.group(3).upper()
        for i, ch in enumerate(letters):
            num = start + i
            if num <= end:
                answers[num] = ch
    return answers


def apply_answers(questions, text):
    answers = parse_answer_key(text)
    for q in questions:
        letter = answers.get(q["number"])
        if letter:
            for opt in q["options"]:
                if opt.startswith(letter + "."):
                    q["correct_answer"] = opt
                    break


# ==========================================================
# 主入口
# ==========================================================

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "message": "缺少文件路径"}, ensure_ascii=False))
        return 1

    # ======================================================
    # 子进程模式：extract.py --ocr-page <pdf> <page_idx> <dpi> <engine>
    # engine: rapidocr（默认，使用 OCR_CONFIG 配置）| tesseract（低内存兜底）
    # OCR 配置通过环境变量传递：OCR_ENGINE, OCR_DET_MODEL, OCR_REC_MODEL, OCR_VERSION
    # ======================================================
    if sys.argv[1] == "--ocr-page" and len(sys.argv) >= 4:
        pdf_path = sys.argv[2]
        page_idx = int(sys.argv[3])
        dpi = int(sys.argv[4]) if len(sys.argv) >= 5 else 120
        engine = sys.argv[5] if len(sys.argv) >= 6 else "rapidocr"

        result = {"success": False, "file": os.path.basename(pdf_path),
                  "page": page_idx, "engine": engine, "text": ""}

        # 从环境变量读取 OCR 配置（由父进程传入）
        ocr_cfg = {}
        for key in ("engine", "det_model", "rec_model", "version"):
            env_key = f"OCR_{key.upper()}"
            val = os.environ.get(env_key, "")
            if val:
                ocr_cfg[key] = val
        # 映射 version 字段
        if "version" in ocr_cfg:
            ocr_cfg["ocr_version"] = ocr_cfg.pop("version")
        ocr_page_config = {**OCR_CONFIG, **ocr_cfg} if ocr_cfg else None

        # 渲染单页 → 临时 PNG
        tmp_png = None
        try:
            import pymupdf as fitz
            doc = fitz.open(pdf_path)
            if page_idx < 0 or page_idx >= len(doc):
                result["message"] = "page out of range: %d/%d" % (page_idx, len(doc))
                print(json.dumps(result, ensure_ascii=False))
                return 1
            page = doc[page_idx]
            zoom = dpi / 72.0
            mat = fitz.Matrix(zoom, zoom)
            pm = page.get_pixmap(matrix=mat, alpha=False)
            tmp_fd, tmp_png = tempfile.mkstemp(suffix=".png", prefix="ocr_p%d_" % page_idx)
            os.close(tmp_fd)
            pm.save(tmp_png)
            doc.close()
            del pm, page, mat, doc
            import gc; gc.collect()
        except Exception as exc:
            result["message"] = "render failed: %s" % exc
            print(json.dumps(result, ensure_ascii=False))
            return 1

        # OCR 引擎
        ocr_text = ""
        if engine == "rapidocr":
            # RapidOCR 路径（使用 OCR_CONFIG 配置，内存约 38MB，精度 100%）
            try:
                ocr_text = _ocr_rapidocr_images([tmp_png], config=ocr_page_config) or ""
            except Exception:
                ocr_text = ""
        elif engine == "tesseract":
            # Tesseract CLI 路径（内存约 50MB，精度较低但稳定）
            try:
                ocr_text = _ocr_tesseract_cli(tmp_png) or ""
            except Exception:
                ocr_text = ""

        if tmp_png:
            try:
                os.unlink(tmp_png)
            except Exception:
                pass

        result["success"] = bool(ocr_text.strip())
        result["text"] = ocr_text

        try:
            sys.stdout.buffer.write(json.dumps(result, ensure_ascii=False).encode("utf-8"))
            sys.stdout.buffer.write(b"\n")
        except Exception:
            print(json.dumps(result, ensure_ascii=False))
        return 0 if result["success"] else 1

    path = sys.argv[1]
    do_parse = "--parse" in sys.argv
    force_ocr = "--ocr" in sys.argv

    result = {
        "success": False,
        "file": os.path.basename(path),
        "type": file_type(path),
        "method": "",
        "text": "",
    }

    if not os.path.isfile(path):
        result["message"] = "文件不存在"
        print(json.dumps(result, ensure_ascii=False))
        return 1

    ftype = result["type"]
    try:
        if ftype == "pdf":
            text, method = extract_pdf(path, force_ocr=force_ocr)
        elif ftype == "doc":
            text, method = extract_doc(path)
        elif ftype == "docx":
            text, method = extract_docx(path)
        elif ftype == "txt":
            text, method = extract_txt(path), "txt"
        else:
            result["message"] = "不支持的文件类型: %s" % ftype
            print(json.dumps(result, ensure_ascii=False))
            return 1
    except Exception as exc:
        result["message"] = "提取失败: %s" % exc
        print(json.dumps(result, ensure_ascii=False))
        return 1

    if not text.strip():
        result["message"] = "未能提取到任何文本"
        print(json.dumps(result, ensure_ascii=False))
        return 1

    result["success"] = True
    result["method"] = method
    result["text"] = text

    if do_parse:
        questions = parse_choice_questions(text)
        apply_answers(questions, text)
        result["questions"] = questions

    # 确保 stdout 在中文 Windows 下以 UTF-8 输出，避免 shell_exec 的 PHP 再收到问号
    try:
        sys.stdout.buffer.write(json.dumps(result, ensure_ascii=False).encode("utf-8"))
        sys.stdout.buffer.write(b"\n")
    except Exception:
        print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
